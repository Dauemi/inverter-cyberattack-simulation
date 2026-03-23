from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


ASSETS = [
    ("Resilience Curve", r"C:\Users\gemim\.cursor\projects\c-Users-gemim-OneDrive-Bureau-M1-cours-Data-engineer-MSC-1-AI-Semestre-2-Pratice-Based-Project-Inverter-cyberattack-simulation\assets\c__Users_gemim_AppData_Roaming_Cursor_User_workspaceStorage_ac14ed6699ff7621edd7a3bbbb5e0d41_images_image-d613acde-27b9-4205-bc38-d3207b7d1946.png"),
    ("Risk Matrix", r"C:\Users\gemim\.cursor\projects\c-Users-gemim-OneDrive-Bureau-M1-cours-Data-engineer-MSC-1-AI-Semestre-2-Pratice-Based-Project-Inverter-cyberattack-simulation\assets\c__Users_gemim_AppData_Roaming_Cursor_User_workspaceStorage_ac14ed6699ff7621edd7a3bbbb5e0d41_images_image-b37122fc-b79a-4f5b-a895-21063ce71372.png"),
]


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)


def add_images_section(doc: Document) -> None:
    doc.add_heading("Screenshots of Key Graphs", level=1)
    fig_num = 1
    for title, path in ASSETS:
        if not os.path.exists(path):
            continue
        doc.add_picture(path, width=Inches(6.0))
        cap = doc.add_paragraph(f"Figure {fig_num}: {title}")
        cap.runs[0].italic = True
        doc.add_paragraph(" ")
        fig_num += 1


def add_table_of_contents(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    # Static (filled) TOC for read-only use
    toc_items = [
        "1. Project Objective",
        "2. Data Sources",
        "3. Simulation Steps (Code Pipeline)",
        "4. Dashboard Overview",
        "5. Dashboard Tabs and Graphs",
        "6. How to Run (Quick Guide)",
        "7. Limitations",
        "8. Deployment (Vercel)",
        "Screenshots of Key Graphs",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_paragraph(" ")


def add_list_of_figures(doc: Document) -> None:
    doc.add_heading("List of Figures", level=1)
    # Static (filled) list of figures
    fig_num = 1
    for title, path in ASSETS:
        if not os.path.exists(path):
            continue
        doc.add_paragraph(f"Figure {fig_num}: {title}")
        fig_num += 1
    doc.add_paragraph(" ")


def build_full_doc(out_path: str) -> None:
    doc = Document()
    add_title(doc, "2ASICYA Project Overview")
    doc.add_paragraph("Solar Inverter Cyber Attack Simulation & Dashboard")
    doc.add_paragraph(" ")
    add_table_of_contents(doc)
    add_list_of_figures(doc)

    doc.add_heading("1. Project Objective", level=1)
    doc.add_paragraph(
        "Model, simulate, and evaluate the impact of coordinated cyberattacks on distributed solar PV "
        "inverters in a French distribution grid using pandapower and synthetic datasets."
    )

    doc.add_heading("2. Data Sources", level=1)
    doc.add_paragraph().add_run("A) France Sprint3 Synthetic Grid").bold = True
    doc.add_paragraph("Files:")
    for f in [
        "fr_grid_buses.csv",
        "fr_grid_lines.csv",
        "fr_grid_loads.csv",
        "fr_grid_pv_generators.csv",
        "fr_inverter_parameters.csv",
        "fr_load_profile_15min.csv",
        "fr_attack_scenarios_S1_S5.csv",
    ]:
        doc.add_paragraph(f"- {f}")

    doc.add_paragraph().add_run("B) Kaggle Datasets (for dashboard enrichment)").bold = True
    doc.add_paragraph("- Solar Power Generation Data: https://www.kaggle.com/datasets/anikannal/solar-power-generation-data")
    doc.add_paragraph("- Smart Grid Intrusion Detection: https://www.kaggle.com/datasets/hussainsheikh03/smart-grid-intrusion-detection-dataset")

    doc.add_heading("3. Simulation Steps (Code Pipeline)", level=1)
    for s in [
        "Load grid topology (buses, lines, loads, PV generators).",
        "Load time-series profile (15-min load multiplier).",
        "Apply PV daily shape (sunrise → sunset).",
        "Inject attack scenario (S1–S5) at 12:00.",
        "Run AC power flow at each timestep with pandapower.",
        "Collect outputs: min/max voltage, line loading, PV output.",
        "Monte Carlo loop for variability across runs.",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("4. Dashboard Overview", level=1)
    doc.add_paragraph(
        "The dashboard is a static HTML (2ASICYA_Dashboard.html) that reads a JSON object (const D = {...}) "
        "generated from the combined datasets. A RUN + RELOAD button refreshes the data when hosted via a local server."
    )

    doc.add_heading("5. Dashboard Tabs and Graphs", level=1)
    doc.add_heading("5.1 Overview Tab", level=2)
    for s in [
        "KPI tiles: Grid Frequency, Installed PV, Grid Buses, Peak Load, Grid Status.",
        "PV Production & Irradiance (24h) – line + area chart.",
        "Load Profile & Temperature (24h) – line + area chart.",
        "Voltage by Bus (p.u.) – multi-line plot.",
        "Line Loading (%) – multi-line plot.",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("5.2 Attack Simulation Tab", level=2)
    for s in [
        "Scenario selector S1–S5 (with affected power levels).",
        "KPIs: Lost Power, Compromised Inverters, Attack Duration, Frequency Deviation.",
        "PV Production: Normal vs Attack (area comparison).",
        "Frequency Impact (Hz) with UFLS thresholds.",
        "Voltage Impact (p.u. deviation) area chart.",
        "Scenario Comparison — Lost Power (bar chart).",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("5.3 IDS & Security Tab", level=2)
    for s in [
        "KPIs: total alerts, detected, missed, detection rate.",
        "Alerts by Hour (bar chart).",
        "Severity Breakdown (donut/pie chart).",
        "Alert Types by MITRE ATT&CK ICS (horizontal bar).",
        "Latest Alerts feed (table-like list).",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("5.4 Grid & Topology Tab", level=2)
    for s in [
        "Network topology visualization (33-bus).",
        "Inverter parameters table.",
        "PV generators table.",
        "Daily profile by PV generator (stacked area).",
        "Load distribution by bus (bar chart).",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("5.5 Risk Matrix Tab", level=2)
    for s in [
        "Risk matrix (Likelihood vs Impact) scatter plot.",
        "Risk Score by Scenario (bar chart).",
        "Risk Scenario Details table.",
        "Risk scores are computed from voltage, frequency, and line loading metrics.",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("5.6 Resilience Curve Tab", level=2)
    for s in [
        "Resilience Differentiation Curve (Collapse Probability vs Attack Magnitude).",
        "Two lines: No Mitigation vs Mitigation Enabled.",
        "Interpretation: mitigation curve should be lower to indicate improved resilience.",
    ]:
        doc.add_paragraph(f"- {s}")

    add_images_section(doc)

    doc.add_heading("6. How to Run (Quick Guide)", level=1)
    for s in [
        "Install dependencies: pip install -r requirements.txt",
        "Run simulation: python -m src.main",
        "Smoke test scenarios: python scripts/smoke_scenarios.py",
        "Build combined dashboard data: python dashboard/build_dashboard_data.py --html 2ASICYA_Dashboard.html --out dashboard/data/combined_dashboard.json --france-dir data/france_sprint3 --kaggle-dir dashboard/data --inject",
        "Serve dashboard locally: python -m http.server 8000",
        "Open: http://localhost:8000/2ASICYA_Dashboard.html",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("7. Limitations", level=1)
    for s in [
        "Frequency model is simplified (no inertia or FCR).",
        "Local feeder scale is extrapolated to GW scenarios.",
        "IDS is conceptual; not a real-time implementation.",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("8. Deployment (Vercel)", level=1)
    for s in [
        "Vercel serves 2ASICYA_Dashboard.html at the root.",
        "The index.html redirect provides a clean URL.",
        "Production URL: https://inverter-cyberattack-simulation.vercel.app/",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.save(out_path)


def build_short_doc(out_path: str) -> None:
    doc = Document()
    add_title(doc, "2ASICYA — Short Project Summary")
    doc.add_paragraph("1–2 page overview for presentation.")
    doc.add_paragraph(" ")
    add_table_of_contents(doc)
    add_list_of_figures(doc)

    doc.add_heading("Project Goal", level=1)
    doc.add_paragraph(
        "Simulate coordinated cyberattacks on PV inverters in a French distribution grid and visualize "
        "their impact on voltage, line loading, and frequency."
    )

    doc.add_heading("Data", level=1)
    doc.add_paragraph("France Sprint3 synthetic grid + Kaggle solar/IDS datasets (dashboard enrichment).")

    doc.add_heading("Method (High Level)", level=1)
    for s in [
        "Pandapower AC power‑flow time‑series (15‑min resolution).",
        "Attack scenarios S1–S5 injected at 12:00.",
        "Monte Carlo variability for robustness.",
        "Dashboard refresh via combined JSON data.",
    ]:
        doc.add_paragraph(f"- {s}")

    doc.add_heading("Key Visuals", level=1)
    doc.add_paragraph("Overview, Attack Simulation, IDS, Grid Topology, Risk Matrix, Resilience Curve.")

    add_images_section(doc)

    doc.add_heading("Limitations", level=1)
    doc.add_paragraph("- Simplified frequency model (no inertia/FCR).")
    doc.add_paragraph("- Local feeder extrapolated to GW‑scale scenarios.")
    doc.add_paragraph("- IDS is conceptual (not real‑time).")

    doc.save(out_path)


if __name__ == "__main__":
    out_dir = os.path.join(os.getcwd(), "docs")
    os.makedirs(out_dir, exist_ok=True)
    build_full_doc(os.path.join(out_dir, "2ASICYA_Project_Overview.docx"))
    print(os.path.join(out_dir, "2ASICYA_Project_Overview.docx"))
